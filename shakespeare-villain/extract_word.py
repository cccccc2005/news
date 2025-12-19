"""
从Word文档(.docx)中提取莎士比亚三个剧本的文本，并识别反派角色的台词
"""
import os
import re
import pandas as pd
from pathlib import Path

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("请安装 python-docx: pip install python-docx")


def extract_text_from_docx(docx_path):
    """从Word文档提取所有文本（包括段落和表格）"""
    if not HAS_DOCX:
        raise ImportError("请安装 python-docx: pip install python-docx")
    
    doc = Document(docx_path)
    text_parts = []
    
    # 提取段落文本
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)
    
    # 提取表格文本
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                # 表格行用制表符分隔，或者每行单独一行
                text_parts.append("\t".join(row_texts))
    
    return "\n".join(text_parts)


def identify_play_sections(full_text):
    """
    识别三个剧本的章节：
    - 哈姆雷特（Hamlet）
    - 麦克白（Macbeth）
    - 奥赛罗（Othello）
    """
    plays = {
        "哈姆雷特": [],
        "麦克白": [],
        "奥赛罗": []
    }
    
    lines = full_text.split('\n')
    current_play = None
    current_text = []
    
    # 更灵活的匹配模式
    play_keywords = {
        "哈姆雷特": ["哈姆雷特", "Hamlet", "HAMLET"],
        "麦克白": ["麦克白", "Macbeth", "MACBETH"],
        "奥赛罗": ["奥赛罗", "Othello", "OTHELLO"]
    }
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        # 检查是否是剧本标题（可能是单独一行，或包含在行中）
        for play_name, keywords in play_keywords.items():
            for keyword in keywords:
                # 检查整行是否主要是标题
                if keyword in line_stripped:
                    # 如果这一行看起来像标题（短且包含关键词）
                    if len(line_stripped) < 50 or line_stripped.startswith(keyword):
                        # 保存之前的剧本
                        if current_play and current_text:
                            plays[current_play].extend(current_text)
                        current_play = play_name
                        current_text = []
                        break
            if current_play:
                break
        
        if current_play:
            current_text.append(line)
    
    # 保存最后一个剧本
    if current_play and current_text:
        plays[current_play].extend(current_text)
    
    return plays


def extract_character_lines(text, character_name, play_name):
    """
    从文本中提取特定角色的台词
    朱生豪译本格式通常是：角色名 + 冒号/空格 + 台词
    """
    lines = []
    text_lines = text.split('\n')
    
    # 匹配模式1：角色名：台词（冒号可能是中文或英文）
    pattern1 = re.compile(
        rf"^\s*{re.escape(character_name)}\s*[：:]\s*(.+)$"
    )
    
    # 匹配模式2：角色名单独一行，下一行是台词
    # 匹配模式3：角色名（台词）或 角色名 台词（无标点）
    
    for i, line in enumerate(text_lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        # 模式1：角色名：台词
        match = pattern1.match(line_stripped)
        if match:
            dialogue = match.group(1).strip()
            if dialogue and len(dialogue) > 2:  # 至少3个字符
                lines.append({
                    "play": play_name,
                    "character": character_name,
                    "text": dialogue,
                    "raw_line": line_stripped
                })
                continue
        
        # 模式2：角色名单独一行，下一行是台词
        if line_stripped == character_name and i + 1 < len(text_lines):
            next_line = text_lines[i + 1].strip()
            # 下一行不是空，不是另一个角色名，不是幕次场次
            if (next_line and 
                len(next_line) > 2 and
                not re.match(r"^第[一二三四五六七八九十]+", next_line) and
                not re.match(r"^[一二三四五六七八九十]+", next_line)):
                lines.append({
                    "play": play_name,
                    "character": character_name,
                    "text": next_line,
                    "raw_line": f"{line_stripped}\n{next_line}"
                })
                continue
        
        # 模式3：行首是角色名，后面直接跟台词（无冒号）
        if line_stripped.startswith(character_name):
            remaining = line_stripped[len(character_name):].strip()
            # 去掉可能的标点
            remaining = re.sub(r"^[：:\s]+", "", remaining)
            if remaining and len(remaining) > 2:
                lines.append({
                    "play": play_name,
                    "character": character_name,
                    "text": remaining,
                    "raw_line": line_stripped
                })
    
    return lines


def parse_act_scene_from_text(text_lines):
    """
    尝试识别幕次和场次
    返回每行对应的act和scene
    """
    act_scene_pattern = re.compile(
        r"(第[一二三四五六七八九十]+幕)\s*(第[一二三四五六七八九十]+场)?",
        re.IGNORECASE
    )
    
    act_map = {
        "一": "1", "二": "2", "三": "3", "四": "4", "五": "5",
        "六": "6", "七": "7", "八": "8", "九": "9", "十": "10"
    }
    
    current_act = "1"
    current_scene = "1"
    act_scene_list = []
    
    for line in text_lines:
        match = act_scene_pattern.search(line)
        if match:
            act_str = match.group(1)
            scene_str = match.group(2) if match.group(2) else None
            
            # 转换中文数字
            for cn, num in act_map.items():
                if cn in act_str:
                    current_act = num
                    break
            
            if scene_str:
                for cn, num in act_map.items():
                    if cn in scene_str:
                        current_scene = num
                        break
        
        act_scene_list.append({
            "act": current_act,
            "scene": current_scene
        })
    
    return act_scene_list


def identify_play_from_filename(filename):
    """从文件名识别剧本"""
    filename_lower = filename.lower()
    if "哈姆雷特" in filename or "hamlet" in filename_lower:
        return "哈姆雷特"
    elif "麦克白" in filename or "macbeth" in filename_lower:
        return "麦克白"
    elif "奥赛罗" in filename or "othello" in filename_lower:
        return "奥赛罗"
    return None


def identify_play_from_content(text):
    """从文本内容识别剧本（检查前2000字符）"""
    preview = text[:2000].lower()
    play_keywords = {
        "哈姆雷特": ["哈姆雷特", "hamlet", "丹麦王子", "丹麦"],
        "麦克白": ["麦克白", "macbeth", "苏格兰", "邓肯"],
        "奥赛罗": ["奥赛罗", "othello", "威尼斯", "苔丝狄蒙娜"]
    }
    
    scores = {}
    for play_name, keywords in play_keywords.items():
        score = sum(1 for kw in keywords if kw.lower() in preview)
        scores[play_name] = score
    
    if scores and max(scores.values()) > 0:
        return max(scores.items(), key=lambda x: x[1])[0]
    return None


def main():
    base_dir = Path(__file__).parent.parent
    
    # 查找所有Word文档
    word_files = list(base_dir.glob("*.docx")) + list(base_dir.glob("*.doc"))
    
    if not word_files:
        print(f"错误: 在 {base_dir} 中找不到Word文档(.docx或.doc)")
        print("请将三个Word文档（哈姆雷特、麦克白、奥赛罗）放在项目根目录")
        return
    
    print(f"找到 {len(word_files)} 个Word文档:")
    for f in word_files:
        print(f"  - {f.name}")
    
    # 三个剧本的配置
    # 注意：哈姆雷特中克劳狄斯的台词可能以"国王"或"国  王"标识
    play_configs = {
        "哈姆雷特": {"character": "克劳狄斯", "alt_names": ["国王", "国  王"], "file": None},
        "麦克白": {"character": "麦克白", "alt_names": [], "file": None},
        "奥赛罗": {"character": "伊阿古", "alt_names": [], "file": None}
    }
    
    # 为每个剧本匹配对应的Word文档
    print("\n正在识别每个文档对应的剧本...")
    
    # 先按文件名精确匹配
    for word_file in word_files:
        # 跳过临时文件和演讲稿文件
        if word_file.name.startswith('.~') or '新闻稿子' in word_file.name or '演讲稿' in word_file.name:
            print(f"  跳过: {word_file.name} (临时文件或演讲稿)")
            continue
        
        play_name = identify_play_from_filename(word_file.name)
        
        if play_name and play_name in play_configs:
            if play_configs[play_name]["file"] is None:
                play_configs[play_name]["file"] = word_file
                print(f"  ✓ {word_file.name} -> {play_name} (文件名匹配)")
            else:
                print(f"  警告: {play_name} 已有匹配文件，跳过 {word_file.name}")
    
    # 如果还有未匹配的剧本，尝试从内容识别
    for play_name, config in play_configs.items():
        if config["file"] is None:
            for word_file in word_files:
                # 跳过临时文件和演讲稿文件
                if word_file.name.startswith('.~') or '新闻稿子' in word_file.name or '演讲稿' in word_file.name:
                    continue
                
                # 跳过已匹配的文件
                if any(c["file"] == word_file for c in play_configs.values() if c["file"]):
                    continue
                
                try:
                    text = extract_text_from_docx(str(word_file))
                    identified = identify_play_from_content(text)
                    if identified == play_name:
                        play_configs[play_name]["file"] = word_file
                        print(f"  ✓ {word_file.name} -> {play_name} (内容识别)")
                        break
                except Exception as e:
                    continue
    
    # 检查是否所有剧本都有对应的文件
    missing_plays = [name for name, config in play_configs.items() if config["file"] is None]
    if missing_plays:
        print(f"\n警告: 以下剧本未找到对应文件: {', '.join(missing_plays)}")
        print("请确保文件名包含剧本名称（如：哈姆雷特.docx、麦克白.docx、奥赛罗.docx）")
    
    # 保存原始文本用于调试
    output_dir = base_dir / "shakespeare-villain" / "output"
    output_dir.mkdir(exist_ok=True)
    
    # 提取每个剧本的台词
    all_lines = []
    
    for play_name, config in play_configs.items():
        if config["file"] is None:
            print(f"\n跳过 {play_name}（未找到对应文件）")
            continue
        
        word_file = config["file"]
        character_name = config["character"]
        
        print(f"\n{'='*60}")
        print(f"处理: {play_name} ({word_file.name})")
        print(f"提取角色: {character_name}")
        print(f"{'='*60}")
        
        try:
            full_text = extract_text_from_docx(str(word_file))
            print(f"文本长度: {len(full_text)} 字符")
            
            # 保存原始文本
            raw_text_path = output_dir / f"raw_text_{play_name}.txt"
            with open(raw_text_path, "w", encoding="utf-8") as f:
                f.write(full_text)
            print(f"原始文本已保存: {raw_text_path}")
            
            # 提取角色台词
            print(f"\n正在提取 {character_name} 的台词...")
            character_lines = extract_character_lines(full_text, character_name, play_name)
            
            # 如果主要角色名提取失败，尝试使用替代名称（如"国王"代表"克劳狄斯"）
            if len(character_lines) < 10 and "alt_names" in config:
                for alt_name in config.get("alt_names", []):
                    print(f"  尝试使用替代名称: {alt_name}")
                    alt_lines = extract_character_lines(full_text, alt_name, play_name)
                    if len(alt_lines) > len(character_lines):
                        character_lines = alt_lines
                        print(f"  使用 {alt_name} 找到 {len(character_lines)} 条台词")
            
            # 统一角色名（如果使用了替代名称，改为标准名称）
            for line in character_lines:
                line["character"] = character_name
            
            print(f"找到 {len(character_lines)} 条台词")
            
            if len(character_lines) > 0:
                print(f"\n前5条示例:")
                for i, line in enumerate(character_lines[:5]):
                    print(f"  {i+1}. {line['text'][:60]}...")
            
            all_lines.extend(character_lines)
            
        except Exception as e:
            print(f"处理 {play_name} 时出错: {e}")
            import traceback
            traceback.print_exc()
            continue
    
    # 转换为DataFrame
    if all_lines:
        df = pd.DataFrame(all_lines)
        
        # 为每条台词分配幕次场次（简化处理，默认值）
        # 后续可以根据文本内容改进，自动识别幕次场次
        df["act"] = "1"
        df["scene"] = "1"
        df["to"] = ""  # 对话对象，需要人工标注或后续改进
        df["is_interrupt"] = 0
        
        # 重新排列列顺序
        df = df[["play", "act", "scene", "character", "to", "is_interrupt", "text"]]
        
        # 保存CSV
        csv_path = output_dir / "villain_lines.csv"
        df.to_csv(csv_path, index=False, encoding="utf-8-sig")
        
        print(f"\n{'='*60}")
        print(f"✓ 提取完成！")
        print(f"{'='*60}")
        print(f"总共提取 {len(df)} 条台词")
        print(f"已保存到: {csv_path}")
        print(f"\n各角色台词统计:")
        stats = df.groupby(["play", "character"]).size()
        for (play, char), count in stats.items():
            print(f"  {play} - {char}: {count} 条")
        
        print(f"\n数据预览（前10条）:")
        print(df.head(10).to_string())
        
    else:
        print("\n" + "="*60)
        print("警告: 未找到任何台词")
        print("="*60)
        print("可能的原因:")
        print("1. Word文档格式与预期不符")
        print("2. 角色名称不匹配（请检查是否是'克劳狄斯'、'麦克白'、'伊阿古'）")
        print("3. 台词格式不是'角色名：台词'")
        print("\n建议:")
        print("1. 查看 output/raw_text_*.txt 检查文本提取是否正确")
        print("2. 如果提取正确但未识别到台词，可能需要手动调整提取规则")
        print("3. 确保Word文档文件名包含剧本名称（如：哈姆雷特.docx）")


if __name__ == "__main__":
    main()

